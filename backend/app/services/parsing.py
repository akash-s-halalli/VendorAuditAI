"""Document parsing service for text extraction."""

import io
from dataclasses import dataclass, field
from datetime import datetime
from typing import ClassVar

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class ParsedPage:
    """Represents a parsed page from a document."""

    page_number: int
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Represents a fully parsed document."""

    pages: list[ParsedPage]
    total_pages: int
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        """Get the full text of the document."""
        return "\n\n".join(page.text for page in self.pages if page.text)

    @property
    def has_tables(self) -> bool:
        """Check if document contains any tables."""
        return any(page.tables for page in self.pages)


class PDFParser:
    """Parser for PDF documents using PyMuPDF."""

    @staticmethod
    def parse(content: bytes) -> ParsedDocument:
        """Parse a PDF document and extract text and tables.

        Args:
            content: PDF file bytes

        Returns:
            ParsedDocument with extracted content

        Raises:
            ValueError: If PDF cannot be parsed
        """
        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception as e:
            raise ValueError(f"Failed to open PDF: {e!s}") from e

        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]

            # Extract text
            text = page.get_text("text")

            # Extract tables (basic approach using blocks)
            tables = PDFParser._extract_tables(page)

            pages.append(ParsedPage(
                page_number=page_num + 1,
                text=text.strip(),
                tables=tables,
            ))

        # Extract metadata
        metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "producer": doc.metadata.get("producer", ""),
            "creation_date": PDFParser._parse_pdf_date(doc.metadata.get("creationDate", "")),
            "modification_date": PDFParser._parse_pdf_date(doc.metadata.get("modDate", "")),
        }

        doc.close()

        return ParsedDocument(
            pages=pages,
            total_pages=len(pages),
            metadata=metadata,
        )

    @staticmethod
    def _extract_tables(page: fitz.Page) -> list[list[list[str]]]:
        """Extract tables from a PDF page.

        This is a simplified table extraction. For production,
        consider using more sophisticated table detection.

        Args:
            page: PyMuPDF page object

        Returns:
            List of tables, each table is a list of rows
        """
        tables = []

        # Use PyMuPDF's table finder if available (v1.23+)
        try:
            page_tables = page.find_tables()
            for table in page_tables:
                extracted = table.extract()
                if extracted:
                    tables.append(extracted)
        except AttributeError:
            # Fallback for older PyMuPDF versions - no table extraction
            pass

        return tables

    @staticmethod
    def _parse_pdf_date(date_str: str) -> str | None:
        """Parse PDF date format (D:YYYYMMDDHHmmSS) to ISO format.

        Args:
            date_str: PDF date string

        Returns:
            ISO format date string or None
        """
        if not date_str:
            return None

        try:
            # Remove D: prefix if present
            if date_str.startswith("D:"):
                date_str = date_str[2:]

            # Parse basic format YYYYMMDDHHMMSS
            if len(date_str) >= 14:
                dt = datetime.strptime(date_str[:14], "%Y%m%d%H%M%S")
                return dt.isoformat()
            elif len(date_str) >= 8:
                dt = datetime.strptime(date_str[:8], "%Y%m%d")
                return dt.isoformat()
        except ValueError:
            pass

        return None


class DOCXParser:
    """Parser for DOCX documents using python-docx."""

    @staticmethod
    def parse(content: bytes) -> ParsedDocument:
        """Parse a DOCX document and extract text and tables.

        Args:
            content: DOCX file bytes

        Returns:
            ParsedDocument with extracted content

        Raises:
            ValueError: If DOCX cannot be parsed
        """
        try:
            doc = DocxDocument(io.BytesIO(content))
        except Exception as e:
            raise ValueError(f"Failed to open DOCX: {e!s}") from e

        # DOCX doesn't have pages in the same way as PDF
        # We'll treat the whole document as one "page" with sections
        text_parts = []
        tables = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        # Extract metadata from core properties
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "creator": core_props.author or "",
            "creation_date": core_props.created.isoformat() if core_props.created else None,
            "modification_date": core_props.modified.isoformat() if core_props.modified else None,
        }

        # Create single page with all content
        page = ParsedPage(
            page_number=1,
            text="\n\n".join(text_parts),
            tables=tables,
        )

        return ParsedDocument(
            pages=[page],
            total_pages=1,  # DOCX doesn't have physical pages until rendered
            metadata=metadata,
        )


class DocumentParser:
    """Main document parser that delegates to specific parsers."""

    PARSERS: ClassVar[dict[str, type[PDFParser] | type[DOCXParser]]] = {
        "application/pdf": PDFParser,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXParser,
        "application/msword": DOCXParser,  # Older .doc format - may not work perfectly
    }

    @classmethod
    def parse(cls, content: bytes, mime_type: str) -> ParsedDocument:
        """Parse a document based on its MIME type.

        Args:
            content: File bytes
            mime_type: MIME type of the document

        Returns:
            ParsedDocument with extracted content

        Raises:
            ValueError: If MIME type is not supported or parsing fails
        """
        parser_class = cls.PARSERS.get(mime_type)
        if not parser_class:
            raise ValueError(f"Unsupported document type: {mime_type}")

        return parser_class.parse(content)

    @classmethod
    def supported_types(cls) -> list[str]:
        """Get list of supported MIME types."""
        return list(cls.PARSERS.keys())
